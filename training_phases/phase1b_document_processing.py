#!/usr/bin/env python3
"""
Phase 1b: Document Processing & Embeddings (90 min)
Day 1 - Understanding RAG: Ingesting and parsing enterprise data

Learning Objectives:
- How to extract text from various document formats
- Understanding embeddings and semantic similarity
- Processing and chunking documents for AI consumption
- Hands-on document ingestion pipeline

This module focuses on the data preparation aspects of RAG without
the complexity of vector databases or retrieval logic.
"""

import os
import re
import hashlib
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path

# Document processing libraries
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ pypdf not available. Install with: pip install pypdf")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available. Install with: pip install python-docx")

# For embeddings demonstration
try:
    import requests
    EMBEDDING_AVAILABLE = True
except ImportError:
    EMBEDDING_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("document-processing")

class DocumentProcessor:
    """
    Core document processing class that handles multiple formats
    and prepares them for AI consumption.
    """
    
    def __init__(self):
        self.processed_documents = []
        self.processing_stats = {
            'total_processed': 0,
            'successful': 0,
            'failed': 0,
            'total_chunks': 0,
            'total_characters': 0
        }
    
    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF files using pypdf.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not PDF_AVAILABLE:
            return {"error": "PDF processing not available", "text": ""}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                metadata = {
                    "num_pages": len(pdf_reader.pages),
                    "file_size": os.path.getsize(file_path),
                    "file_name": os.path.basename(file_path)
                }
                
                # Extract text from all pages
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                # Clean up text
                text = self.clean_text(text)
                
                return {
                    "text": text,
                    "metadata": metadata,
                    "source_type": "pdf",
                    "success": True
                }
                
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "source_type": "pdf",
                "success": False
            }
    
    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word documents.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not DOCX_AVAILABLE:
            return {"error": "DOCX processing not available", "text": ""}
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "file_size": os.path.getsize(file_path),
                "file_name": os.path.basename(file_path)
            }
            
            # Extract text from all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Clean up text
            text = self.clean_text(text)
            
            return {
                "text": text,
                "metadata": metadata,
                "source_type": "docx",
                "success": True
            }
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "source_type": "docx",
                "success": False
            }
    
    def extract_text_from_txt(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from plain text files.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "file_size": os.path.getsize(file_path),
                "file_name": os.path.basename(file_path),
                "encoding": "utf-8"
            }
            
            # Clean up text
            text = self.clean_text(text)
            
            return {
                "text": text,
                "metadata": metadata,
                "source_type": "txt",
                "success": True
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text: {e}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "source_type": "txt",
                "success": False
            }
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\-\:\;\(\)]', ' ', text)
        
        # Remove page markers
        text = re.sub(r'--- Page \d+ ---', '', text)
        
        # Strip and return
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks for better AI processing.
        
        Args:
            text: Text to chunk
            chunk_size: Target size for each chunk (in characters)
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks with metadata
        """
        chunks = []
        words = text.split()
        
        current_chunk = []
        current_length = 0
        chunk_id = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            
            if current_length + word_length > chunk_size and current_chunk:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    "id": f"chunk_{chunk_id}",
                    "text": chunk_text,
                    "length": len(chunk_text),
                    "word_count": len(current_chunk),
                    "chunk_index": chunk_id
                })
                
                # Start new chunk with overlap
                overlap_words = current_chunk[-overlap//10:] if overlap > 0 else []
                current_chunk = overlap_words + [word]
                current_length = sum(len(w) + 1 for w in current_chunk)
                chunk_id += 1
            else:
                current_chunk.append(word)
                current_length += word_length
        
        # Add final chunk if there's content
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "text": chunk_text,
                "length": len(chunk_text),
                "word_count": len(current_chunk),
                "chunk_index": chunk_id
            })
        
        return chunks
    
    def process_document(self, file_path: str, chunk_size: int = 500) -> Dict[str, Any]:
        """
        Process a single document: extract text, clean, and chunk.
        
        Args:
            file_path: Path to document
            chunk_size: Size for text chunking
            
        Returns:
            Processed document with chunks and metadata
        """
        start_time = datetime.now()
        file_ext = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            extraction_result = self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            extraction_result = self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            extraction_result = self.extract_text_from_txt(file_path)
        else:
            extraction_result = {
                "text": "",
                "metadata": {"error": f"Unsupported file type: {file_ext}"},
                "source_type": "unknown",
                "success": False
            }
        
        if not extraction_result["success"]:
            self.processing_stats['failed'] += 1
            return extraction_result
        
        # Chunk the text
        chunks = self.chunk_text(extraction_result["text"], chunk_size)
        
        # Create document ID
        doc_id = hashlib.md5(file_path.encode()).hexdigest()[:8]
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Create final document structure
        processed_doc = {
            "id": doc_id,
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "source_type": extraction_result["source_type"],
            "original_text": extraction_result["text"],
            "chunks": chunks,
            "metadata": {
                **extraction_result["metadata"],
                "processing_time": processing_time,
                "chunk_count": len(chunks),
                "total_length": len(extraction_result["text"]),
                "processed_at": start_time.isoformat()
            },
            "success": True
        }
        
        # Update stats
        self.processing_stats['total_processed'] += 1
        self.processing_stats['successful'] += 1
        self.processing_stats['total_chunks'] += len(chunks)
        self.processing_stats['total_characters'] += len(extraction_result["text"])
        
        # Store processed document
        self.processed_documents.append(processed_doc)
        
        logger.info(f"Processed {file_path}: {len(chunks)} chunks, {processing_time:.2f}s")
        
        return processed_doc
    
    def process_directory(self, directory_path: str, chunk_size: int = 500) -> List[Dict[str, Any]]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            chunk_size: Size for text chunking
            
        Returns:
            List of processed documents
        """
        supported_extensions = {'.pdf', '.docx', '.txt'}
        processed_docs = []
        
        directory = Path(directory_path)
        if not directory.exists():
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        for file_path in directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    doc = self.process_document(str(file_path), chunk_size)
                    if doc.get("success", False):
                        processed_docs.append(doc)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    self.processing_stats['failed'] += 1
        
        return processed_docs
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """
        Get statistics about processed documents.
        """
        return {
            **self.processing_stats,
            "average_chunks_per_doc": (
                self.processing_stats['total_chunks'] / max(1, self.processing_stats['successful'])
            ),
            "average_doc_length": (
                self.processing_stats['total_characters'] / max(1, self.processing_stats['successful'])
            )
        }

def simple_embedding_demo():
    """
    Demonstrate basic embedding concepts (requires local embedding service).
    This is a simplified example - in practice you'd use proper embedding models.
    """
    print("🔍 Simple Embedding Concept Demo")
    print("(Note: This is a conceptual demo, not production embeddings)")
    
    # Sample texts
    texts = [
        "How do I return a product?",
        "What is your return policy?",
        "I want to exchange my item",
        "My password is not working",
        "Reset my account password",
        "Shipping takes too long"
    ]
    
    print("\nSample customer queries:")
    for i, text in enumerate(texts, 1):
        print(f"{i}. {text}")
    
    # Simple word-based similarity (not real embeddings)
    print("\nSimple similarity analysis:")
    query = "return policy information"
    query_words = set(query.lower().split())
    
    similarities = []
    for text in texts:
        text_words = set(text.lower().split())
        common_words = query_words & text_words
        similarity = len(common_words) / len(query_words | text_words)
        similarities.append((text, similarity))
    
    # Sort by similarity
    similarities.sort(key=lambda x: x[1], reverse=True)
    
    print(f"\nQuery: '{query}'")
    print("Most similar texts:")
    for text, sim in similarities[:3]:
        print(f"  {sim:.2f}: {text}")

def demo_document_processing():
    """
    Demonstrate document processing capabilities.
    """
    print("=== Phase 1b: Document Processing & Embeddings Demo ===\n")
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # Check available formats
    print("📄 Supported Document Formats:")
    print(f"  ✅ PDF: {'Available' if PDF_AVAILABLE else 'Not available'}")
    print(f"  ✅ DOCX: {'Available' if DOCX_AVAILABLE else 'Not available'}")
    print(f"  ✅ TXT: Available")
    print()
    
    # Process sample documents
    sample_dir = "/Users/developer/capstone/knowledge_base_pdfs"
    
    if os.path.exists(sample_dir):
        print(f"🔄 Processing documents from: {sample_dir}")
        processed_docs = processor.process_directory(sample_dir, chunk_size=300)
        
        print(f"\n📊 Processing Results:")
        stats = processor.get_processing_stats()
        print(f"  • Documents processed: {stats['successful']}/{stats['total_processed']}")
        print(f"  • Total chunks created: {stats['total_chunks']}")
        print(f"  • Average chunks per document: {stats['average_chunks_per_doc']:.1f}")
        print(f"  • Average document length: {stats['average_doc_length']:.0f} characters")
        
        # Show sample chunks
        if processed_docs:
            print(f"\n📝 Sample chunks from '{processed_docs[0]['file_name']}':")
            for i, chunk in enumerate(processed_docs[0]['chunks'][:3]):
                print(f"\nChunk {i+1} ({chunk['length']} chars):")
                print(f"  {chunk['text'][:150]}{'...' if len(chunk['text']) > 150 else ''}")
    else:
        print(f"⚠️ Sample directory not found: {sample_dir}")
        print("Creating a sample text file for demonstration...")
        
        # Create sample content
        sample_content = """
        Company Return Policy
        
        We offer a 30-day return policy for all products. Items must be in original condition
        with original packaging and receipt. Refunds are processed within 5-7 business days.
        
        For returns:
        1. Contact customer service
        2. Receive return authorization
        3. Ship item with prepaid label
        4. Receive confirmation and refund
        
        Shipping Policy
        
        Standard shipping takes 3-5 business days within the US and costs $5.99.
        Express shipping is available in 1-2 days for $15.99.
        Free shipping is offered on orders over $50.
        International shipping takes 7-14 days and costs $19.99.
        """
        
        # Save sample file
        sample_file = "/Users/developer/capstone/training_phases/sample_policy.txt"
        with open(sample_file, 'w') as f:
            f.write(sample_content)
        
        # Process sample file
        doc = processor.process_document(sample_file)
        
        print(f"✅ Processed sample document:")
        print(f"  • Chunks: {len(doc['chunks'])}")
        print(f"  • Total length: {doc['metadata']['total_length']} characters")
        
        print(f"\n📝 Sample chunks:")
        for i, chunk in enumerate(doc['chunks'][:2]):
            print(f"\nChunk {i+1}:")
            print(f"  {chunk['text'][:200]}...")
    
    # Embedding concept demo
    print("\n" + "="*50)
    simple_embedding_demo()

if __name__ == "__main__":
    demo_document_processing()
    
    print("\n🎓 Phase 1b Complete!")
    print("Next: Phase 1c - Vector Database Setup")